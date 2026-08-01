"""阅卷笔录 Word 文档生成器.

按需求定义的七部分结构输出，所有事实/证据均带 见《卷》P x-y 来源引用。
对标律师手工版阅卷笔录的关键结构：
  - 封面标题后带自动目录（TOC 域，Word 中 F9 更新页码）；
  - 供述/证言逐字问答全文渲染（要点 + 全文），标题行含 第N次/卷宗页码/同步录音录像；
  - 程序性文书与书证按待证事实分组，含文号列；流水类书证附逐笔资金流水子表；
  - 阅卷结论前渲染资金勾稽一览表。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

from ..workspace import CaseWorkspace, DocumentaryEvidence, Statement


def _set_cn_font(doc: Document) -> None:
    """设置默认中文字体（宋体）."""
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "SimSun")


def _add_toc(doc: Document) -> None:
    """插入自动目录域（TOC）：Word/WPS 打开后右键“更新域”或按 F9 生成页码."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "（目录：在 Word/WPS 中右键“更新域”或按 F9 生成）"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    doc.add_page_break()


def _cite_text(c) -> str:
    return c.render() if c else "（未标注来源）"


def _cite_full(c) -> str:
    """完整来源引用：见《卷》P起-止（context 说明）。含 JSON 里 source 对象的全部信息。"""
    if not c:
        return "（未标注来源）"
    base = c.render()
    if c.context:
        return f"{base}（{c.context}）"
    return base


def _source_line(doc: Document, c) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"来源：{_cite_full(c)}")
    r.font.size = Pt(9)
    r.italic = True


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if widths:
        for ci, w in enumerate(widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)


def _statement_heading(s: Statement, seq: int) -> str:
    """律师式一行标题：张某 第1次 《证据卷》P27-34（无同步录音录像；刑事拘留后）."""
    parts = [s.person, f"第{seq}次"]
    if s.volume:
        pages = f"P{s.page_start}" if s.page_start == s.page_end else f"P{s.page_start}-{s.page_end}"
        parts.append(f"《{s.volume}》{pages}")
    flags = []
    if s.has_av_recording and s.has_av_recording != "未注明":
        flags.append("有同步录音录像" if s.has_av_recording == "是" else "无同步录音录像")
    elif s.has_av_recording == "未注明":
        flags.append("同步录音录像未注明")
    if s.occasion:
        flags.append(s.occasion)
    if flags:
        parts.append(f"（{'；'.join(flags)}）")
    return " ".join(parts)


def _statement_block(doc: Document, s: Statement, seq: int) -> None:
    """渲染一条笔录（供述/证言）：一行式标题 + 元信息 + 要点 + 逐字全文 + 来源."""
    doc.add_heading(_statement_heading(s, seq), level=3)
    meta = doc.add_paragraph()
    meta.add_run(
        f"笔录时间：{s.record_time or '未注明'}    办案人员：{s.investigators or '未注明'}    "
        f"办案地点：{s.location or '未注明'}"
    ).font.size = Pt(10)
    if s.charged_fact_ref:
        p = doc.add_paragraph()
        p.add_run(f"对应指控事实：{s.charged_fact_ref}").font.size = Pt(10)
    if s.content_summary:
        p = doc.add_paragraph()
        p.add_run("要点：").bold = True
        p.add_run(s.content_summary)
    if s.full_text:
        p = doc.add_paragraph()
        p.add_run("笔录全文：").bold = True
        for line in s.full_text.splitlines():
            doc.add_paragraph(line.strip()) if line.strip() else None
    else:
        p = doc.add_paragraph()
        r = p.add_run("（⚠ 未登记逐字全文，仅有要点——建议复核后补登）")
        r.font.size = Pt(9)
        r.italic = True
    _source_line(doc, s.source)


def _person_seq(statements: list[Statement]) -> list[int]:
    """为每条笔录计算“同一人第几次”序号（按列表顺序）."""
    seen: dict[str, int] = {}
    seqs = []
    for s in statements:
        seen[s.person] = seen.get(s.person, 0) + 1
        seqs.append(seen[s.person])
    return seqs


def _group_by_fact(items, default_group: str) -> dict[str, list]:
    """按 fact_group 分组（保持登记顺序）；空组名归入默认组."""
    groups: dict[str, list] = {}
    for it in items:
        groups.setdefault((it.fact_group or "").strip() or default_group, []).append(it)
    return groups


def _transactions_block(doc: Document, e: DocumentaryEvidence) -> None:
    """渲染书证下的逐笔资金流水子表."""
    if not e.transactions:
        return
    p = doc.add_paragraph()
    p.add_run(f"资金流水（{e.name}，共 {len(e.transactions)} 笔）：").bold = True
    _add_table(
        doc,
        ["序号", "日期", "付款方", "收款方", "金额", "账户", "页码", "备注"],
        [
            [str(i), tx.date or "—", tx.payer or "—", tx.payee or "—", tx.amount or "—",
             tx.account or "—", f"P{tx.page}" if tx.page else "—", tx.note or "—"]
            for i, tx in enumerate(e.transactions, 1)
        ],
        widths=[1.0, 2.2, 2.2, 2.2, 2.2, 3.2, 1.2, 2.3],
    )


def generate_review_notes(ws: CaseWorkspace, out_dir: Path) -> Path:
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _set_cn_font(doc)

    case = ws.case_name or "案件"
    title = doc.add_heading(f"{case} 阅卷笔录", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _add_toc(doc)

    # ---- 案件基本信息表 ----
    doc.add_heading("案件基本信息表", level=1)
    # 完整渲染，不截断：指控事实全文（若过长则单独成段，避免表格单元格压缩丢失文字）
    fact_full = ws.indictment.full_text_summary or "—"
    key_evidence = "；".join(
        f"{e.name}({_cite_text(e.source_ref)})" for e in ws.documentary_evidence
    ) or "（见下文第七部分）"
    _add_table(
        doc,
        ["项目", "内容"],
        [
            ["1. 代理主体", ws.defendant or "—"],
            ["2. 涉嫌罪名", ws.charge or "—"],
            ["3. 起诉书/起诉意见书指控事实", fact_full],
            ["4. 涉案金额", ws.total_amount or "—"],
            ["5. 卷宗数量", f"{ws.volume_count} 册（{'; '.join('《'+v+'》' for v in ws.known_volumes())}）" if ws.known_volumes() else f"{ws.volume_count} 册"],
            ["6. 关键证据索引", key_evidence],
        ],
        widths=[4.5, 12.0],
    )

    # ---- 一、当事人基本情况 ----
    doc.add_heading("一、当事人基本情况", level=1)
    p = ws.party
    if p.name:
        lines = [
            f"姓名：{p.name}",
            f"性别：{p.gender or '—'}    民族：{p.ethnicity or '—'}    出生：{p.birth or '—'}",
            f"籍贯：{p.native_place or '—'}    身份证号：{p.id_no or '—'}    文化程度：{p.education or '—'}",
            f"职业：{p.occupation or '—'}    住址：{p.address or '—'}",
        ]
        for ln in lines:
            doc.add_paragraph(ln)
        if p.position or p.appointment_history:
            doc.add_heading("任职情况（职务犯罪）", level=3)
            if p.position:
                doc.add_paragraph(f"职务：{p.position}")
            for ap in p.appointment_history:
                doc.add_paragraph(ap, style="List Bullet")
        _source_line(doc, p.source)
    else:
        doc.add_paragraph("（卷宗中未提取到当事人基本情况，待补充。）")

    # ---- 二、起诉书、起诉意见书内容 ----
    doc.add_heading("二、起诉书、起诉意见书内容", level=1)
    ind = ws.indictment
    if ind.doc_type:
        doc.add_paragraph(
            f"文书类型：{ind.doc_type}    制作机关：{ind.issuer or '—'}    落款日期：{ind.issue_date or '—'}"
        )
        doc.add_paragraph(f"被告人：{ind.defendant or '—'}    涉嫌罪名：{ind.charge or '—'}    涉案金额：{ind.total_amount or '—'}")
        if ind.legal_basis:
            doc.add_paragraph("适用法律条文：")
            for lb in ind.legal_basis:
                doc.add_paragraph(lb, style="List Bullet")
        if ind.sentencing_circumstances:
            doc.add_paragraph("量刑情节：")
            for sc in ind.sentencing_circumstances:
                doc.add_paragraph(sc, style="List Bullet")
        doc.add_heading("指控事实（原文摘录）", level=3)
        doc.add_paragraph(ind.full_text_summary or "（无）")
        _source_line(doc, ind.source)
        if ind.facts:
            doc.add_heading("指控事实分笔", level=3)
            _add_table(
                doc,
                ["序号", "事实概述", "金额", "时间", "来源"],
                [
                    [str(f.index), f.description, f.amount or "—", f.time_period or "—", _cite_text(f.source)]
                    for f in ind.facts
                ],
                widths=[1.2, 8.5, 2.2, 2.5, 2.5],
            )
    else:
        doc.add_paragraph("（卷宗中未发现起诉书/起诉意见书，待补充。）")

    # ---- 三、被告人供述和辩解 ----
    doc.add_heading("三、被告人的供述和辩解", level=1)
    if ws.defendant_statements:
        groups: dict[str, list[Statement]] = {}
        for s in ws.defendant_statements:
            groups.setdefault(s.charged_fact_ref or "（未对应具体事实）", []).append(s)
        for key, stmts in groups.items():
            doc.add_heading(f"对应：{key}", level=2)
            for s, seq in zip(stmts, _person_seq(stmts)):
                _statement_block(doc, s, seq)
    else:
        doc.add_paragraph("（未提取到被告人供述笔录。）")

    # ---- 四、同案人员供述和辩解 ----
    doc.add_heading("四、同案人员的供述和辩解", level=1)
    if ws.codefendant_statements:
        for s, seq in zip(ws.codefendant_statements, _person_seq(ws.codefendant_statements)):
            _statement_block(doc, s, seq)
    else:
        doc.add_paragraph("（本案无同案人员，或卷宗中未提取到同案人供述。）")

    # ---- 五、证人证言 ----
    doc.add_heading("五、证人证言", level=1)
    if ws.witness_statements:
        for s, seq in zip(ws.witness_statements, _person_seq(ws.witness_statements)):
            _statement_block(doc, s, seq)
    else:
        doc.add_paragraph("（卷宗中未提取到证人证言。）")

    # ---- 六、程序性文书 ----
    doc.add_heading("六、程序性文书", level=1)
    if ws.procedural_docs:
        for group, docs_ in _group_by_fact(ws.procedural_docs, "程序性文书").items():
            if len(ws.procedural_docs) != len(docs_) or group != "程序性文书":
                doc.add_heading(group, level=2)
            _add_table(
                doc,
                ["序号", "文书类型", "文号", "时间", "地点", "主要内容", "来源"],
                [
                    [str(i + 1), d.doc_type, d.doc_no or "—", d.time or "—", d.location or "—",
                     d.content_summary, _cite_text(d.source)]
                    for i, d in enumerate(docs_)
                ],
                widths=[1.0, 2.4, 3.0, 2.0, 2.0, 4.0, 2.1],
            )
    else:
        doc.add_paragraph("（未提取到程序性文书。）")

    # ---- 七、书证（按待证事实分组）----
    doc.add_heading("七、书证（客观证据，按待证事实分组）", level=1)
    if ws.documentary_evidence:
        for group, evs in _group_by_fact(ws.documentary_evidence, "其他书证").items():
            doc.add_heading(group, level=2)
            _add_table(
                doc,
                ["序号", "文件名称", "编号", "形成时间", "卷宗页码", "来源/制作主体", "主要内容", "来源引用"],
                [
                    [str(i + 1), e.name, e.doc_no or "—", e.time or "—",
                     f"P{e.page_start}-{e.page_end}" if e.page_start else "—",
                     e.source or "—", e.content_summary, _cite_text(e.source_ref)]
                    for i, e in enumerate(evs)
                ],
                widths=[0.9, 2.4, 2.0, 1.8, 1.7, 2.0, 4.0, 2.0],
            )
            for e in evs:
                _transactions_block(doc, e)
    else:
        doc.add_paragraph("（未提取到书证。）")

    # ---- 资金勾稽一览 ----
    f = ws.funds
    funds_rows = [
        ("报案/投资人陈述金额合计", f.reported_amount),
        ("合同/协议金额合计", f.contract_amount),
        ("起诉书指控金额", f.charged_amount),
        ("已返还/返利金额", f.returned_amount),
        ("违法所得（工资/提成）", f.illegal_income),
        ("已退赔金额", f.restitution),
    ]
    if any(v for _, v in funds_rows) or f.note:
        doc.add_heading("资金勾稽一览", level=1)
        _add_table(
            doc,
            ["项目", "金额/说明"],
            [[k, v or "—"] for k, v in funds_rows] + ([["勾稽说明", f.note]] if f.note else []),
            widths=[6.0, 10.5],
        )

    # ---- 阅卷结论 ----
    doc.add_heading("阅卷结论", level=1)
    c = ws.conclusions
    doc.add_heading("（一）已查明核心事实", level=3)
    for x in c.core_facts:
        doc.add_paragraph(x, style="List Bullet")
    if not c.core_facts:
        doc.add_paragraph("（待补充）")
    doc.add_heading("（二）证据链条", level=3)
    for x in c.evidence_chain:
        doc.add_paragraph(x, style="List Bullet")
    if not c.evidence_chain:
        doc.add_paragraph("（待补充）")
    doc.add_heading("（三）证据矛盾点", level=3)
    for x in c.contradictions:
        doc.add_paragraph(x, style="List Bullet")
    if not c.contradictions:
        doc.add_paragraph("（未发现明显矛盾）")
    doc.add_heading("（四）待核查疑点", level=3)
    for x in c.doubts:
        doc.add_paragraph(x, style="List Bullet")
    if not c.doubts:
        doc.add_paragraph("（暂无）")
    # raw：结论合成员的完整原文（若有），完整保留，不截断
    if c.raw:
        doc.add_heading("（五）结论原文", level=3)
        doc.add_paragraph(c.raw)

    # ---- 阅卷目录附表 ----
    doc.add_heading("附：阅卷目录", level=1)
    _add_table(
        doc,
        ["卷宗名称", "页码", "所含文件", "文书类型", "笔录时间", "备注"],
        [
            [e.volume_name, e.page_range, e.file_name, e.doc_type, e.record_time, e.note]
            for e in ws.catalog
        ]
        or [["—"] * 6],
        widths=[3.0, 1.8, 4.0, 2.5, 2.5, 2.5],
    )

    # 免责声明
    doc.add_paragraph()
    disc = doc.add_paragraph()
    r = disc.add_run(
        "说明：本笔录由阅卷 Agent 基于本地卷宗材料自动梳理生成，所有引用均标注来源卷宗及页码，"
        "仅供后续辩护意见参考，不构成正式辩护策略或出庭意见。"
    )
    r.font.size = Pt(9)
    r.italic = True

    safe = case.replace("/", "_").replace(" ", "")
    path = out_dir / f"{safe}_阅卷笔录.docx"
    doc.save(str(path))
    return path
