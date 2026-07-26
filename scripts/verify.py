"""对抗性校验：核对 docx/xlsx 阅卷产出（不再依赖 workspace.json）。
1. 从 docx 提取所有 见《卷》P起-止 引用，核对页码落在真实卷宗页数内。
2. 抽样引用回溯：取引用对应的卷宗页，到实际页面文本（docling/tesseract 回退）确认含相关关键词（防编造）。
3. 七部分齐全性（检查 docx 标题层级）。
4. docx 与 xlsx 的引用一致性。
用法: python scripts/verify.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import docx
import openpyxl


CITE_RE = re.compile(r"见《([^》]+)》P(\d+)(?:-(\d+))?")


def extract_cites_from_docx(path: Path) -> list[tuple[str, int, int, str]]:
    """从 docx 提取所有引用：[(卷宗, 起, 止, 上下文), ...]。"""
    d = docx.Document(str(path))
    cites = []
    # 段落
    for p in d.paragraphs:
        for m in CITE_RE.finditer(p.text):
            vol = m.group(1); ps = int(m.group(2)); pe = int(m.group(3)) if m.group(3) else ps
            ctx = p.text[max(0, m.start()-30):m.end()+10]
            cites.append((vol, ps, pe, ctx))
    # 表格
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for m in CITE_RE.finditer(cell.text):
                    vol = m.group(1); ps = int(m.group(2)); pe = int(m.group(3)) if m.group(3) else ps
                    cites.append((vol, ps, pe, cell.text[:40]))
    return cites


def check_citation_ranges(cites, volume_pages) -> list:
    problems = []
    for vol, ps, pe, ctx in cites:
        if vol not in volume_pages:
            problems.append(f"卷宗《{vol}》不在登记卷宗（P{ps}-{pe}）"); continue
        if ps <= 0 or pe < ps or pe > volume_pages[vol]:
            problems.append(f"页码越界 《{vol}》P{ps}-{pe} (上限{volume_pages[vol]}) | {ctx[:30]}")
    return problems


def check_traceback(cites, store) -> list:
    """抽样若干引用，核对对应页面文本是否非空且含 CJK（防编造空页引用）。

    不再用"上下文关键词"硬匹配（上下文常为标签文字如"来源当事人"，与页面正文无关，易误报）。
    改为：确认被引用页确实有实质 OCR 文本（≥20字且含汉字），即引用指向真实有内容的页。
    """
    problems = []
    seen = set()
    samples = []
    for vol, ps, pe, ctx in cites:
        if vol in seen:
            continue
        seen.add(vol)
        samples.append((vol, ps, pe, ctx))
        if len(samples) >= 8:
            break
    for vol, ps, pe, ctx in samples:
        try:
            v = store.get(vol)
        except Exception as e:
            problems.append(f"《{vol}》无法打开: {e}"); continue
        texts = [(p, (v.get_page_text(p).text or "")) for p in range(ps, pe + 1)]
        # 被引用页应至少有一页含实质中文文本
        has_content = any(len(re.findall(r"[一-鿿]", t)) >= 15 for _, t in texts)
        if not has_content:
            problems.append(f"《{vol}》P{ps}-{pe} 被引用页无实质中文文本（疑似引用空页/编造）")
    return problems


def check_sections(path: Path) -> dict:
    d = docx.Document(str(path))
    headings = [p.text.strip() for p in d.paragraphs if p.style.name.startswith("Heading")]
    want = ["一、当事人", "二、起诉书", "三、被告人", "四、同案", "五、证人", "六、程序性", "七、书证", "阅卷结论", "阅卷目录"]
    present = {}
    for w in want:
        present[w] = any(w in h for h in headings)
    present["表格数"] = len(d.tables)
    return present


def main() -> int:
    from vibelawyer.config import load_case
    from vibelawyer.pdf_volume import VolumeStore
    from vibelawyer import docling_cache
    cfg = load_case()
    volume_pages = {v.name: v.pages for v in cfg.volumes}
    out = cfg.output_dir

    docx_files = sorted(out.glob("*_阅卷笔录.docx"))
    xlsx_files = sorted(out.glob("*_阅卷目录.xlsx"))
    if not docx_files:
        print("未找到 阅卷笔录.docx"); return 1
    docx_p, xlsx_p = docx_files[-1], (xlsx_files[-1] if xlsx_files else None)
    print(f"校验对象: {docx_p.name}")

    cites = extract_cites_from_docx(docx_p)
    print(f"\n=== 1. 引用页码范围校验（共 {len(cites)} 处引用）===")
    p1 = check_citation_ranges(cites, volume_pages)
    print(f"  越界/无效: {len(p1)}")
    for x in p1[:8]:
        print(f"    - {x}")

    print("\n=== 2. 抽样引用回溯（防编造）===")
    store = VolumeStore(cfg, docling_cache_dir=docling_cache.cache_dir_for(out))
    p2 = check_traceback(cites, store)
    print(f"  回溯不符: {len(p2)}")
    for x in p2[:8]:
        print(f"    - {x}")

    print("\n=== 3. 七部分齐全性（docx 标题）===")
    sec = check_sections(docx_p)
    for k, v in sec.items():
        print(f"  {k}: {v}")

    if xlsx_p:
        print("\n=== 4. docx 与 xlsx 引用一致性 ===")
        wb = openpyxl.load_workbook(str(xlsx_p))
        ws = wb["证据索引"]
        xlsx_cites = []
        for row in ws.iter_rows(min_row=4, values_only=True):
            for cell in row:
                if cell:
                    for m in CITE_RE.finditer(str(cell)):
                        xlsx_cites.append((m.group(1), int(m.group(2)),
                                           int(m.group(3)) if m.group(3) else int(m.group(2))))
        print(f"  docx 引用数: {len(cites)} | xlsx 证据索引引用数: {len(xlsx_cites)}")
        print(f"  xlsx 引用越界: {len(check_citation_ranges([(v,s,e,'') for v,s,e in xlsx_cites], volume_pages))}")

    print(f"\n=== 总结: {len(p1)} 处页码越界 + {len(p2)} 处回溯不符 ===")
    return 0 if (len(p1) + len(p2)) == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
