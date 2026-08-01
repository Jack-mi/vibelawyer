"""渲染冒烟：合成带全字段的 CaseWorkspace，渲染 docx+xlsx 并断言关键结构.

不调用 LLM、不依赖卷宗 PDF；用于验证 generators/ 升级后结构对标律师版。
用法: python scripts/smoke_render.py
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document

from vibelawyer.generators.docx_notes import generate_review_notes
from vibelawyer.generators.xlsx_catalog import generate_catalog_xlsx
from vibelawyer.workspace import (
    CaseWorkspace,
    CatalogEntry,
    ChargedFact,
    Citation,
    Conclusions,
    DocumentaryEvidence,
    FundsSummary,
    IndictmentInfo,
    PartyInfo,
    ProceduralDoc,
    Statement,
    Transaction,
)


def _cite(volume: str, s: int, e: int, ctx: str = "") -> Citation:
    return Citation(volume=volume, page_start=s, page_end=e, context=ctx)


def build_synthetic_workspace() -> CaseWorkspace:
    ws = CaseWorkspace()
    ws.register_volume("证据卷", 60)
    ws.register_volume("诉讼卷", 40)
    ws.case_name = "张某非法吸收公众存款案"
    ws.defendant = "张某"
    ws.charge = "非法吸收公众存款罪"
    ws.total_amount = "人民币2000万元"
    ws.volume_count = 2

    ws.party = PartyInfo(
        name="张某", gender="男", ethnicity="汉", birth="1975年3月", native_place="北京",
        id_no="110xxx1990030", education="大学本科", occupation="公司法定代表人",
        position="董事长", appointment_history=["2018年起任甲公司董事长"], address="北京市朝阳区",
        other=[], source=_cite("诉讼卷", 1, 3, "当事人基本情况"),
    )

    ws.indictment = IndictmentInfo(
        doc_type="起诉书", issuer="北京市人民检察院分院", issue_date="2024年9月1日",
        defendant="张某", charge="非法吸收公众存款罪", total_amount="人民币2000万元",
        facts=[
            ChargedFact(index=1, description="2022年1月至2023年6月，以投资甲公司项目为名吸收资金",
                         amount="人民币1200万元", time_period="2022.1-2023.6",
                         source=_cite("诉讼卷", 5, 6, "指控事实第1笔")),
            ChargedFact(index=2, description="2023年7月至2023年12月，虚构理财产品吸收资金",
                         amount="人民币800万元", time_period="2023.7-2023.12",
                         source=_cite("诉讼卷", 6, 7, "指控事实第2笔")),
        ],
        legal_basis=["《刑法》第一百七十六条"],
        sentencing_circumstances=["自首", "部分退赔"],
        full_text_summary="被告人张某自2022年1月起，未经依法批准，以投资甲公司项目为名...",
        source=_cite("诉讼卷", 4, 7, "起诉书"),
    )

    ws.defendant_statements = [
        Statement(
            person="张某", role="defendant", volume="证据卷", page_start=10, page_end=18,
            record_time="2024年3月1日", investigators="王某、李某", location="某区看守所",
            has_av_recording="是", occasion="刑事拘留后", charged_fact_ref="第1笔",
            content_summary="供认吸收资金过程，对金额有部分辩解",
            full_text="问：你何时开始吸收资金的？\n答：2022年1月。\n问：吸收了多少？\n答：约1200万。",
            source=_cite("证据卷", 10, 18, "被告人第1次供述"),
        ),
        Statement(
            person="张某", role="defendant", volume="证据卷", page_start=27, page_end=34,
            record_time="2024年4月2日", investigators="王某、李某", location="某区看守所",
            has_av_recording="未注明", occasion="逮捕后", charged_fact_ref="第2笔",
            content_summary="对虚构理财产品部分翻供",
            full_text="问：理财产品是否真实？\n答：部分是虚构的。\n问：具体金额？\n答：约800万。",
            source=_cite("证据卷", 27, 34, "被告人第2次供述"),
        ),
    ]
    ws.codefendant_statements = [
        Statement(
            person="王某", role="codefendant", volume="证据卷", page_start=40, page_end=45,
            record_time="2024年3月15日", investigators="王某、李某", location="某区看守所",
            has_av_recording="否", occasion="取保候审期间", charged_fact_ref="第1笔",
            content_summary="确认参与吸收资金",
            full_text="问：你是否参与？\n答：是，我负责签约。",
            source=_cite("证据卷", 40, 45, "同案人供述"),
        ),
    ]
    ws.witness_statements = [
        Statement(
            person="赵某", role="witness", volume="证据卷", page_start=50, page_end=55,
            record_time="2024年3月20日", investigators="王某、李某", location="某区派出所",
            has_av_recording="未注明", occasion="证人询问", charged_fact_ref="第1笔;第2笔",
            content_summary="投资受害，说明被吸收资金经过",
            full_text="问：你怎么认识张某的？\n答：朋友介绍投资。\n问：投了多少？\n答：30万。",
            source=_cite("证据卷", 50, 55, "证人证言"),
        ),
    ]

    ws.procedural_docs = [
        ProceduralDoc(doc_type="刑事拘留证", volume="诉讼卷", page_start=8, page_end=9,
                      doc_no="京公朝拘字[2024]008号", time="2024年3月1日",
                      location="某区分局", fact_group="强制措施",
                      content_summary="对张某采取刑事拘留", source=_cite("诉讼卷", 8, 9, "拘留证")),
        ProceduralDoc(doc_type="逮捕证", volume="诉讼卷", page_start=12, page_end=13,
                      doc_no="京公朝捕字[2024]012号", time="2024年4月8日",
                      location="某区分局", fact_group="强制措施",
                      content_summary="对张某执行逮捕", source=_cite("诉讼卷", 12, 13, "逮捕证")),
        ProceduralDoc(doc_type="搜查证", volume="诉讼卷", page_start=15, page_end=15,
                      doc_no="京公朝搜字[2024]015号", time="2024年3月2日",
                      location="张某办公室", fact_group="侦查措施",
                      content_summary="搜查办公场所", source=_cite("诉讼卷", 15, 15, "搜查证")),
    ]

    # 书证：含文号、待证事实分组、资金流水
    bank = DocumentaryEvidence(
        name="银行流水（甲公司账户）", volume="证据卷", page_start=60, page_end=60,
        doc_no="流水编号BK2024", time="2022-2023", source="银行出具", fact_group="资金往来",
        content_summary="甲公司对公账户全部收支流水", transactions=[],
        source_ref=_cite("证据卷", 60, 60, "银行流水"),
    )
    ws.documentary_evidence = [bank]
    # 逐笔流水挂载到书证
    for d, payer, payee, amt, pg in [
        ("2022年3月1日", "赵某", "甲公司", "人民币30万元", 60),
        ("2022年5月10日", "钱某", "甲公司", "人民币50万元", 60),
        ("2023年1月15日", "甲公司", "孙某", "人民币5万元", 60),
    ]:
        ok, _ = ws.attach_transaction("银行流水（甲公司账户）",
                                      Transaction(date=d, payer=payer, payee=payee,
                                                  amount=amt, account="招行6225xxxx", note="投资款/返利", page=pg))
        assert ok, f"流水挂载失败: {d}"

    ws.catalog = [
        CatalogEntry(volume_name="证据卷", page_range="1-60", file_name="讯问笔录/证言/银行流水",
                     doc_type="言词证据+书证", record_time="2024年", note="第1册"),
        CatalogEntry(volume_name="诉讼卷", page_range="1-40", file_name="起诉书/强制措施文书",
                     doc_type="程序性文书", record_time="2024年", note="第2册"),
    ]

    ws.conclusions = Conclusions(
        core_facts=["张某主导甲公司吸收资金2000万元", "部分资金用于返利"],
        evidence_chain=["起诉书→被告人供述→银行流水印证金额"],
        contradictions=["第2次供述对虚构理财产品部分翻供，与银行流水不符"],
        doubts=["未退赔金额具体去向待查"],
        raw="综合全卷，张某行为符合非法吸收公众存款罪构成要件。",
    )

    ws.funds = FundsSummary(
        reported_amount="人民币2000万元", contract_amount="人民币2000万元",
        charged_amount="人民币2000万元", returned_amount="人民币200万元",
        illegal_income="人民币1500万元", restitution="人民币200万元",
        note="报案、合同、指控三项口径一致；已返还200万计入退赔。",
    )
    return ws


def assert_docx(path: Path, ws: CaseWorkspace) -> None:
    doc = Document(str(path))
    xml = doc.element.xml
    paragraphs = "\n".join(p.text for p in doc.paragraphs)
    tables = doc.tables
    table_text = "\n".join(c.text for t in tables for r in t.rows for c in r.cells)

    checks: list[tuple[str, bool, str]] = []

    def chk(name: str, ok: bool, detail: str = "") -> None:
        checks.append((name, ok, detail))

    # 1. TOC 域
    chk("TOC自动目录域", 'TOC' in xml and 'fldSimple' in xml)
    # 2. 供述一行式标题（含 第N次 + 卷宗页码 + 同步录音录像 标记）
    chk("供述一行式标题(第1次/页码/录音录像)",
        "张某 第1次" in paragraphs and "证据卷" in paragraphs and ("有同步录音录像" in paragraphs or "无同步录音录像" in paragraphs))
    # 3. 逐字全文渲染
    chk("笔录全文逐字渲染", "问：你何时开始吸收资金的" in paragraphs and "答：2022年1月" in paragraphs)
    # 4. 程序性文书按 fact_group 分组
    chk("程序文书分组(强制措施/侦查措施)",
        "强制措施" in paragraphs and "侦查措施" in paragraphs)
    # 5. 文号列（拘留证/逮捕证文号出现在表格中）
    chk("程序文书文号列", "京公朝拘字[2024]008号" in table_text and "京公朝捕字[2024]012号" in table_text)
    # 6. 书证文号 + 资金流水子表
    chk("书证文号", "流水编号BK2024" in table_text)
    chk("资金流水子表(逐笔)", "人民币30万元" in table_text and "人民币50万元" in table_text and "赵某" in table_text)
    # 7. 量刑情节
    chk("量刑情节渲染", "量刑情节" in paragraphs and "自首" in paragraphs)
    # 8. 资金勾稽一览表
    chk("资金勾稽一览表", "资金勾稽一览" in paragraphs and "人民币2000万元" in table_text and "报案合计" in table_text or "reported" in table_text or "报案" in table_text)
    chk("资金勾稽-退赔项", "人民币200万元" in table_text)
    # 9. 来源引用标注
    chk("来源引用标注", "见《证据卷》" in paragraphs or "见《诉讼卷》" in paragraphs)

    print("\n===== docx 结构断言 =====")
    all_ok = True
    for name, ok, detail in checks:
        flag = "✓" if ok else "✗"
        print(f"  [{flag}] {name}")
        if not ok:
            all_ok = False
    if not all_ok:
        raise SystemExit("docx 结构断言失败，见上")


def assert_xlsx(path: Path) -> None:
    from openpyxl import load_workbook
    wb = load_workbook(str(path))
    sheets = wb.sheetnames
    print(f"\n  xlsx sheets: {sheets}")
    expected = ["分卷总览", "阅卷目录", "案件信息", "证据索引"]
    for s in expected:
        if s not in sheets:
            raise SystemExit(f"xlsx 缺失 sheet: {s}（现有 {sheets}）")
    # 分卷总览：每卷一行 + P起-止:文件名
    ov = wb["分卷总览"]
    ov_text = "\n".join(str(c.value) for r in ov.iter_rows() for c in r if c.value)
    print(f"  分卷总览首单元: {ov_text[:120]!r}")
    if "证据卷" not in ov_text or "P1-60" not in ov_text:
        raise SystemExit("分卷总览缺少卷宗或 P起-止:文件名 提要")
    # 案件信息：资金勾稽 + 流水笔数
    info = wb["案件信息"]
    info_text = "\n".join(str(c.value) for r in info.iter_rows() for c in r if c.value)
    if "资金勾稽-报案合计" not in info_text or "资金流水笔数" not in info_text:
        raise SystemExit("案件信息缺少资金勾稽或流水笔数行")
    print(f"  xlsx size: {path.stat().st_size} bytes")


def main() -> None:
    out = ROOT / "output" / "smoke"
    out.mkdir(parents=True, exist_ok=True)
    ws = build_synthetic_workspace()
    print(f"合成工作区: 卷数={ws.volume_count} 供述={len(ws.defendant_statements)} "
          f"书证={len(ws.documentary_evidence)} 流水={sum(len(e.transactions) for e in ws.documentary_evidence)}")

    docx_path = generate_review_notes(ws, out)
    print(f"docx → {docx_path} ({docx_path.stat().st_size} bytes)")
    assert_docx(docx_path, ws)

    xlsx_path = generate_catalog_xlsx(ws, out)
    print(f"xlsx → {xlsx_path} ({xlsx_path.stat().st_size} bytes)")
    assert_xlsx(xlsx_path)

    print("\n✓ 渲染冒烟通过：docx/xlsx 关键结构对标律师版。")


if __name__ == "__main__":
    main()
